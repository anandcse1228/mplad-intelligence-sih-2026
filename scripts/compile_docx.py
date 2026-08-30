import os
import glob
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def convert_md_to_docx(md_path, docx_path):
    if not HAS_DOCX:
        print("python-docx package not installed. Skipping DOCX compilation.")
        return
    doc = docx.Document()
    doc.add_heading('MPLAD Intelligence Documentation', 0)
    with open(md_path, 'r', encoding='utf-8') as f:
        for line in f:
            line_str = line.strip()
            if line_str.startswith('# '):
                doc.add_heading(line_str[2:], level=1)
            elif line_str.startswith('## '):
                doc.add_heading(line_str[3:], level=2)
            elif line_str.startswith('### '):
                doc.add_heading(line_str[4:], level=3)
            elif line_str.startswith('- ') or line_str.startswith('* '):
                doc.add_paragraph(line_str[2:], style='List Bullet')
            elif line_str:
                doc.add_paragraph(line_str)
    doc.save(docx_path)
    print(f"Compiled {os.path.basename(md_path)} -> {os.path.basename(docx_path)}")

if __name__ == '__main__':
    docs_dir = 'docs' if os.path.exists('docs') else '../docs'
    md_files = glob.glob(os.path.join(docs_dir, '*.md'))
    for md in md_files:
        docx_out = os.path.splitext(md)[0] + '.docx'
        convert_md_to_docx(md, docx_out)
